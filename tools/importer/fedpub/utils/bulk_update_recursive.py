import os
import json
# import requests
from docx import Document

# Define the directory to search for docx files
directory = "bulk_update_dir"

# Load the remote JSON file
json_url = 'path-to-json.json'

with open(json_url, "r") as json_file:
    data = json.load(json_file)

# Loop through the directory recursively
for root, dirs, files in os.walk(directory):
    for file in files:
        # Only process docx files
        if file.endswith(".docx") and not file.startswith("~$"):
            # Load the docx file and get the revision history
            doc = Document(os.path.join(root, file))
            revisions = doc.core_properties.revision

            # Loop through the data and perform replacements
            for item in data["data"]:
                url_parts = item["URL"].split("/")
                doc_name = url_parts[-1].replace(".html", "")
                # if os.path.splittext(file)[0] doc_name == file[:-5]:
                # print(file[:-5])
                if doc_name == file[:-5]:
                    print("Matched: " + doc_name)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if "<w:drawing" in cell._element.xml:
                                    continue
                                # Perform the find/replace
                                cell.text = cell.text.replace("CaaS Primary Tag", item["primaryTag"])
                                cell.text = cell.text.replace("CaaS Tags", item["caasTags"])

                    # Find the Topics row in a table and replace the text in the next cell
                    for table in doc.tables:
                        for row in table.rows:
                            if row.cells[0].text.strip() == "Topics":
                                row.cells[1].text = item["category"]

            # Save the modified document with the original revision history
            doc.core_properties.revision = revisions
            doc.save(os.path.join(root, file))